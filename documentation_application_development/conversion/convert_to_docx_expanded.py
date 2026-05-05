"""Generate the EXPANDED Word document for CS6004ES Application Development Coursework 2.

This script combines the expanded main report, the individual contribution report, and
the user manual into a single DOCX with London Met / ESOFT style formatting:

- A4 page size, 1 inch margins
- Times New Roman 12 pt body, justified, 1.5 line spacing
- Heading 1: 16 pt bold black, Heading 2: 14 pt bold black, Heading 3: 12 pt bold black
- Tables: 10 pt body text, header row bold
- Figure and table captions: 10 pt italic, centred
- Page break between main report, contribution report, and user manual
- Output: documentation_application_development/final_document/Application_Development_Final_Report_EXPANDED.docx

The script does not invent content. It reads the three Markdown files as they are.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = ROOT / "final_document"
OUTPUT = FINAL_DIR / "Application_Development_Final_Report_EXPANDED.docx"

INPUT_FILES = [
    ROOT / "Main_Application_Development_Report.md",
    ROOT / "Expanded_Individual_Contribution.md",
    ROOT / "User_Manual_Application_Development.md",
]

CAPTION_PREFIX_RE = re.compile(
    r"^(?:\[(?:SCREENSHOT REQUIRED|DIAGRAM EXPORT REQUIRED|NEEDS CONFIRMATION|VERIFY)\]\s*)?"
    r"(?:Figure|Table)\s+[A-Za-z0-9\.]+:\s+",
    re.IGNORECASE,
)


def warn(message: str) -> None:
    print(f"WARNING: {message}", file=sys.stderr)


def read_markdown(path: Path) -> str | None:
    if not path.exists():
        warn(f"Missing input file: {path}")
        return None
    return path.read_text(encoding="utf-8")


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [clean_inline_markdown(cell) for cell in line.strip().strip("|").split("|")]


def is_caption_line(text: str) -> bool:
    return bool(CAPTION_PREFIX_RE.match(text.strip()))


def build_document():
    try:
        from docx import Document
    except ImportError as exc:
        warn(f"python-docx is not available: {exc}")
        sys.exit(2)
    return Document()


def configure_document(document) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = document.styles

    def configure_style(style_name: str, size: int, bold: bool, italic: bool = False) -> None:
        if style_name not in styles:
            return
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        try:
            style.font.color.rgb = RGBColor(0, 0, 0)
        except AttributeError:
            pass
        style_element = style.element
        rpr = style_element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:cs"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")

    configure_style("Normal", 12, bold=False)
    configure_style("Heading 1", 16, bold=True)
    configure_style("Heading 2", 14, bold=True)
    configure_style("Heading 3", 12, bold=True)
    configure_style("List Bullet", 12, bold=False)
    configure_style("List Number", 12, bold=False)


def add_caption_paragraph(document, text: str):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(clean_inline_markdown(text))
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    return paragraph


def add_word_field(document, instruction: str, placeholder_text: str) -> None:
    """Insert a real Word field (TOC, TOF, etc.) the marker can update with F9."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instruction} "
    run_instr._r.append(instr_text)

    run_separate = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run_separate._r.append(fld_separate)

    run_placeholder = paragraph.add_run(placeholder_text)
    run_placeholder.font.name = "Times New Roman"
    run_placeholder.font.size = Pt(10)
    run_placeholder.italic = True

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_image_placeholder(document, text: str) -> None:
    """Render a 'Figure X: [INSERT ... HERE]' line as a centred bold placeholder line."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(clean_inline_markdown(text))
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def add_body_paragraph(document, text: str, style: str | None = None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    paragraph = document.add_paragraph(clean_inline_markdown(text), style=style)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    if style is None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_heading(document, text: str, level: int):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
        try:
            run.font.color.rgb = RGBColor(0, 0, 0)
        except AttributeError:
            pass
        run.font.size = Pt({1: 16, 2: 14, 3: 12}.get(level, 12))
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def add_table(document, rows: list[str], pending_caption: str | None) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    parsed_rows = [split_table_row(row) for row in rows if not is_table_separator(row)]
    if not parsed_rows:
        return

    if pending_caption:
        add_caption_paragraph(document, pending_caption)

    max_columns = max(len(row) for row in parsed_rows)
    table = document.add_table(rows=len(parsed_rows), cols=max_columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row in enumerate(parsed_rows):
        for column_index in range(max_columns):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.font.bold = row_index == 0


IMAGE_PLACEHOLDER_RE = re.compile(
    r"^Figure\s+[A-Za-z0-9]+(?:\s*\([^)]+\))?:\s*\[INSERT (?:SCREENSHOT|EXPORTED DIAGRAM IMAGE)\s+HERE\]\s*$"
)

MD_IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def add_markdown_image(document, image_path: Path) -> None:
    """Insert a centred raster image from a Markdown `![](path)` line."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    try:
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.2))
    except Exception as exc:  # pragma: no cover - depends on local image validity
        warn(f"Could not embed image {image_path}: {exc}")
        fail = paragraph.add_run(f"[Image embed failed: {image_path.name}]")
        fail.font.name = "Times New Roman"
        fail.font.size = Pt(11)


def add_markdown_to_docx(document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    index = 0
    pending_table_caption: str | None = None
    expect_caption = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            from docx.shared import Pt

            paragraph = document.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.left_indent = None
            continue

        m_img = MD_IMAGE_RE.match(stripped)
        if m_img:
            raw_rel = m_img.group(2).strip()
            candidate = Path(raw_rel)
            image_path = candidate.resolve() if candidate.is_absolute() else (ROOT / raw_rel).resolve()
            if not image_path.is_file():
                warn(f"Missing image file for markdown embed: {image_path}")
                add_body_paragraph(document, f"[Image file not found: {raw_rel}]")
            else:
                add_markdown_image(document, image_path)
            expect_caption = True
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_rows.append(candidate)
                    index += 1
                else:
                    break
            add_table(document, table_rows, pending_table_caption)
            pending_table_caption = None
            continue

        if stripped == "<<WORD_TOC>>":
            add_word_field(
                document,
                'TOC \\o "1-3" \\h \\z \\u',
                "Right-click and choose Update Field to populate the Table of Contents.",
            )
        elif stripped == "<<WORD_LIST_OF_FIGURES>>":
            add_word_field(
                document,
                'TOC \\h \\z \\c "Figure"',
                "Right-click and choose Update Field to populate the List of Figures.",
            )
        elif stripped == "<<WORD_LIST_OF_TABLES>>":
            add_word_field(
                document,
                'TOC \\h \\z \\c "Table"',
                "Right-click and choose Update Field to populate the List of Tables.",
            )
        elif stripped.startswith("# "):
            add_heading(document, stripped[2:], level=1)
        elif stripped.startswith("## "):
            add_heading(document, stripped[3:], level=2)
        elif stripped.startswith("### "):
            add_heading(document, stripped[4:], level=3)
        elif stripped.startswith("- "):
            add_body_paragraph(document, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_body_paragraph(document, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        elif stripped.startswith("Table ") and ":" in stripped[:40]:
            pending_table_caption = stripped
            index += 1
            continue
        elif IMAGE_PLACEHOLDER_RE.match(stripped):
            add_image_placeholder(document, stripped)
            expect_caption = True
            index += 1
            continue
        elif is_caption_line(stripped):
            add_caption_paragraph(document, stripped)
        elif expect_caption and (
            stripped.lower().startswith("this screenshot shows")
            or stripped.lower().startswith("this diagram shows")
        ):
            add_caption_paragraph(document, stripped)
            expect_caption = False
        elif stripped.startswith("---"):
            pass
        else:
            expect_caption = False
            add_body_paragraph(document, stripped)

        index += 1


def insert_page_break(document) -> None:
    from docx.enum.text import WD_BREAK

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def main() -> int:
    sections = []
    for path in INPUT_FILES:
        text = read_markdown(path)
        if text is None:
            return 1
        sections.append((path.name, text))

    document = build_document()
    configure_document(document)

    for index, (name, text) in enumerate(sections):
        if index > 0:
            insert_page_break(document)
        add_markdown_to_docx(document, text)

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
