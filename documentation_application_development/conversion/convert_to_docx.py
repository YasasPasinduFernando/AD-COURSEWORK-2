from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
FINAL_DIR = ROOT / "final_document"
OUTPUT = FINAL_DIR / "Application_Development_Final_Report.docx"

MAIN_REPORT = ROOT / "Main_Application_Development_Report.md"
CONTRIBUTION = ROOT / "Expanded_Individual_Contribution.md"
USER_MANUAL = ROOT / "User_Manual_Application_Development.md"

INPUT_FILES = [
    MAIN_REPORT,
    CONTRIBUTION,
    USER_MANUAL,
]


def warn(message: str) -> None:
    print(f"WARNING: {message}", file=sys.stderr)


def read_markdown(path: Path) -> str | None:
    if not path.exists():
        warn(f"Missing input file: {path}")
        return None
    return path.read_text(encoding="utf-8")


def build_combined_markdown(sections: list[str]) -> str:
    parts = []
    for index, section in enumerate(sections):
        if index > 0:
            parts.append("\n\n\\newpage\n\n")
        parts.append(section)
    return "\n\n".join(parts)


def build_markdown_sections(include_user_manual: bool) -> list[str]:
    paths = INPUT_FILES if include_user_manual else INPUT_FILES[:2]
    sections: list[str] = []
    for path in paths:
        content = read_markdown(path)
        if content:
            sections.append(content)
    return sections


def convert_with_pandoc(markdown_text: str) -> bool:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return False

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    with tempfile.NamedTemporaryFile("w", suffix=".md", delete=False, encoding="utf-8") as temp:
        temp.write(markdown_text)
        temp_path = Path(temp.name)

    try:
        subprocess.run(
            [
                pandoc,
                str(temp_path),
                "-o",
                str(OUTPUT),
                "--from",
                "markdown",
                "--to",
                "docx",
            ],
            check=True,
        )
        return True
    finally:
        try:
            temp_path.unlink()
        except OSError:
            pass


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def split_table_row(line: str) -> list[str]:
    return [clean_inline_markdown(cell) for cell in line.strip().strip("|").split("|")]


def apply_run_font(paragraph, font_name, font_size) -> None:
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size


def add_table_to_docx(document, rows: list[str]) -> None:
    parsed_rows = [split_table_row(row) for row in rows if not is_table_separator(row)]
    if not parsed_rows:
        return

    max_columns = max(len(row) for row in parsed_rows)
    table = document.add_table(rows=len(parsed_rows), cols=max_columns)
    table.style = "Table Grid"

    from docx.shared import Pt

    for row_index, row in enumerate(parsed_rows):
        for column_index in range(max_columns):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
                    run.font.bold = row_index == 0


def add_paragraph(document, text: str, style: str | None = None):
    paragraph = document.add_paragraph(clean_inline_markdown(text), style=style)
    return paragraph


def add_markdown_to_docx(document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_rows = []
            while index < len(lines):
                candidate = lines[index].strip()
                if candidate.startswith("|") and candidate.endswith("|"):
                    table_rows.append(candidate)
                    index += 1
                else:
                    break
            add_table_to_docx(document, table_rows)
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- "):
            add_paragraph(document, stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            add_paragraph(document, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            add_paragraph(document, stripped)
        index += 1


def convert_with_python_docx(markdown_sections: list[str]) -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION_START
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        warn(f"python-docx is not available: {exc}")
        return False

    def set_style_font(style_name: str, size: int, bold: bool = False) -> None:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style_element = style.element
        rpr = style_element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")

    def set_document_section(section) -> None:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    styles = document.styles

    set_document_section(document.sections[0])

    set_style_font("Normal", 12)
    set_style_font("Heading 1", 16, bold=True)
    set_style_font("Heading 2", 14, bold=True)
    set_style_font("Heading 3", 12, bold=True)

    if "List Bullet" in styles:
        set_style_font("List Bullet", 12)
    if "List Number" in styles:
        set_style_font("List Number", 12)

    for index, markdown_text in enumerate(markdown_sections):
        if index > 0:
            new_section = document.add_section(WD_SECTION_START.NEW_PAGE)
            set_document_section(new_section)
        add_markdown_to_docx(document, markdown_text)

    for paragraph in document.paragraphs:
        apply_run_font(paragraph, "Times New Roman", Pt(12))
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for table in document.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)

    document.save(OUTPUT)
    return True


def format_existing_docx(path: Path) -> bool:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
        from docx.shared import RGBColor
    except ImportError as exc:
        warn(f"python-docx is not available for post-formatting: {exc}")
        return False

    document = Document(path)

    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = RGBColor(0, 0, 0)

    for paragraph in document.paragraphs:
        apply_run_font(paragraph, "Times New Roman", Pt(12))
        if paragraph.style.name == "Normal":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = 1.5
            paragraph.paragraph_format.space_after = Pt(6)

    document.save(path)
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert UniManage coursework Markdown files to DOCX.")
    parser.add_argument(
        "--no-user-manual",
        action="store_true",
        help="Exclude the user manual from the combined DOCX.",
    )
    args = parser.parse_args()

    include_user_manual = not args.no_user_manual
    sections = build_markdown_sections(include_user_manual)

    if not sections:
        print("No input Markdown files were found. DOCX was not created.", file=sys.stderr)
        return 1

    combined = build_combined_markdown(sections)

    if convert_with_pandoc(combined):
        format_existing_docx(OUTPUT)
        print(f"Created {OUTPUT}")
        return 0

    if convert_with_python_docx(sections):
        print(f"Created {OUTPUT}")
        return 0

    print("Pandoc and python-docx are not available. Install one option and rerun this script.", file=sys.stderr)
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
